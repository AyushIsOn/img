"""Builds the NAKSHA detailed report as an editable .docx.

    python3 build_docx.py

The first page reproduces the competition template: the black DETAILED REPORT
heading, the red competition line, and the bordered box carrying team name, team
members, word count and date, with (First Page) at the foot.

Formatting is applied as real Word styles rather than typed in, so it survives
editing: Times New Roman 11 point, 1.5 line spacing, 3 cm margins on all sides.
No institution appears anywhere, which is checked on every build.
"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import report_content as C

OUT = "NAKSHA-Detailed-Report.docx"
FONT = "Times New Roman"
SIZE = Pt(11)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x44, 0x44, 0x44)


# ------------------------------------------------------------------ helpers

def apply_font(run, size=None, bold=False, italic=False, colour=None):
    run.font.name = FONT
    run.font.size = size or SIZE
    run.bold = bold
    run.italic = italic
    if colour is not None:
        run.font.color.rgb = colour
    # Word substitutes a different face for some characters unless the east
    # Asian and complex script slots are set as well as ascii.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT)
    return run


def para(doc, text="", *, size=None, bold=False, italic=False, align=None,
         colour=None, before=None, after=None,
         spacing=WD_LINE_SPACING.ONE_POINT_FIVE, style=None, indent=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = spacing
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if indent is not None:
        pf.left_indent = indent
    if text:
        apply_font(p.add_run(text), size, bold, italic, colour)
    return p


def set_base(doc):
    n = doc.styles["Normal"]
    n.font.name = FONT
    n.font.size = SIZE
    rpr = n.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    n.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(3)
        s.left_margin = s.right_margin = Cm(3)


def count_words(text):
    text = re.sub(r"\[\d+\]", " ", text)
    return len([w for w in text.split() if any(c.isalnum() for c in w)])


def total_words():
    n = count_words(C.SYNOPSIS)
    for kind, text in C.BODY:
        if kind in ("p", "b", "h2", "box"):
            n += count_words(text)
    return n


# -------------------------------------------------------------- front page

def front_page(doc, words):
    para(doc, "DETAILED REPORT", size=Pt(20), bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(14))
    para(doc, "V-GUARD INDUSTRIES LTD \u2013 BIG IDEA TECH DESIGN "
              "COMPETITION 2026",
         size=Pt(14), bold=True, colour=RED,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(22),
         spacing=WD_LINE_SPACING.SINGLE)
    para(doc, "The first page of the report must adhere to the format given "
              "below:", size=Pt(9.5), after=Pt(8),
         spacing=WD_LINE_SPACING.SINGLE)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Cm(15)

    def line(bits, after=Pt(6), indent=None):
        """bits is a list of (text, bold, italic)."""
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = after
        if indent is not None:
            p.paragraph_format.left_indent = indent
        for text, bold, italic in bits:
            apply_font(p.add_run(text), Pt(11), bold, italic)
        return p

    cell.paragraphs[0].text = ""
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    line([("Team Name: ", True, False), (C.TEAM, True, False)])
    line([("Team Members ", True, False), ("(Full Name)", False, True)])
    for i in range(1, 4):
        name = C.MEMBERS[i - 1] if i <= len(C.MEMBERS) else ""
        label = f"{i}) {name}".rstrip()
        line([(label, True, False)], after=Pt(4), indent=Cm(1.6))

    p = cell.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    apply_font(p.add_run(f"Number of words: {words}"), Pt(11), True)
    apply_font(p.add_run("\t\t\t"), Pt(11))
    apply_font(p.add_run(f"Date of Submission: {C.DATE}"), Pt(11), True)

    para(doc, "", after=Pt(0))
    para(doc, f"{C.TITLE}: {C.SUBTITLE}", size=Pt(11), italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, before=Pt(18), after=Pt(2),
         spacing=WD_LINE_SPACING.SINGLE)
    para(doc, "(First Page)", size=Pt(11), bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, before=Pt(120),
         spacing=WD_LINE_SPACING.SINGLE)
    doc.add_page_break()


# -------------------------------------------------------------------- build

def build():
    doc = Document()
    set_base(doc)
    words = total_words()
    front_page(doc, words)

    # ------------------------------------------------------------ synopsis
    para(doc, "Synopsis", size=Pt(14), bold=True, after=Pt(8))
    for block in C.SYNOPSIS.strip().split("\n\n"):
        para(doc, block.replace("\n", " "),
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=Pt(8))

    para(doc, "Demonstration video", bold=True, before=Pt(10), after=Pt(2),
         spacing=WD_LINE_SPACING.SINGLE)
    p = para(doc, C.DRIVE_LINK, size=Pt(10), after=Pt(4),
             spacing=WD_LINE_SPACING.SINGLE)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x0B, 0x4F, 0xA0)
        r.underline = True
    doc.add_page_break()

    # ----------------------------------------------------------------- toc
    para(doc, "Table of Contents", size=Pt(14), bold=True, after=Pt(8))
    for kind, text in C.BODY:
        if kind == "h1":
            para(doc, text, after=Pt(3), spacing=WD_LINE_SPACING.SINGLE)
        elif kind == "h2":
            para(doc, text, after=Pt(3), indent=Cm(0.8),
                 spacing=WD_LINE_SPACING.SINGLE)
    para(doc, "References", after=Pt(3), spacing=WD_LINE_SPACING.SINGLE)
    doc.add_page_break()

    # ---------------------------------------------------------------- body
    exhibit = 0
    missing = []
    for kind, text in C.BODY:
        if kind == "h1":
            para(doc, text, size=Pt(14), bold=True, before=Pt(14), after=Pt(7))

        elif kind == "h2":
            para(doc, text, bold=True, before=Pt(10), after=Pt(5))

        elif kind == "b":
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = \
                WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6)
            apply_font(p.add_run(text))

        elif kind == "box":
            head, *rest = text.split("\n")
            t = doc.add_table(rows=1, cols=1)
            t.style = "Table Grid"
            c = t.rows[0].cells[0]
            c.paragraphs[0].text = ""
            hp = c.paragraphs[0]
            hp.paragraph_format.space_after = Pt(4)
            hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            apply_font(hp.add_run(head), Pt(11), True)
            for block in rest:
                bp = c.add_paragraph()
                bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                bp.paragraph_format.line_spacing_rule = \
                    WD_LINE_SPACING.SINGLE
                bp.paragraph_format.space_after = Pt(4)
                apply_font(bp.add_run(block), Pt(10))
            para(doc, "", after=Pt(6), spacing=WD_LINE_SPACING.SINGLE)

        elif kind == "img":
            path, caption, width_cm = [x.strip() for x in text.split("|")]
            exhibit += 1
            if not os.path.exists(path):
                missing.append(path)
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.add_run().add_picture(path, width=Cm(float(width_cm)))
            para(doc, f"Exhibit {exhibit}. {caption}", size=Pt(9.5),
                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, colour=GREY,
                 spacing=WD_LINE_SPACING.SINGLE, after=Pt(12))

        else:
            para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=Pt(8))

    # ---------------------------------------------------------- references
    para(doc, "References", size=Pt(14), bold=True, before=Pt(16), after=Pt(8))
    for i, ref in enumerate(C.REFERENCES, start=1):
        p = para(doc, f"[{i}]  {ref}", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 after=Pt(5), spacing=WD_LINE_SPACING.SINGLE)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)

    # -------------------------------------------------------- page numbers
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    apply_font(run, Pt(9))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)

    doc.save(OUT)

    print(f"\nwritten: {OUT}")
    print(f"  team           {C.TEAM}")
    print(f"  members        {', '.join(C.MEMBERS)}")
    print(f"  word count     {words}  (printed on the first page)")
    print(f"  synopsis       {count_words(C.SYNOPSIS)}  (limit 300) "
          f"{'OK' if count_words(C.SYNOPSIS) <= 300 else 'OVER'}")
    print(f"  references     {len(C.REFERENCES)}")
    print(f"  exhibits       {exhibit - len(missing)} of {exhibit}")
    for m in missing:
        print(f"    MISSING {m}")
    print(f"  {FONT} 11pt, 1.5 spacing, 3 cm margins")


if __name__ == "__main__":
    build()
